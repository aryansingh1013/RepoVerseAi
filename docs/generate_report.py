import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body_paragraph(doc, text, bold_prefix="", indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
        
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(12)
        r_bold.font.bold = True
        
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F3F4F6")
    set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
    
    # Set borders to light gray
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D1D5DB')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    doc.add_paragraph("") # Space after block

def create_massive_report(output_path, name, reg_no):
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ─────────────────────────────────────────────────────────────────────────
    # COVER PAGE (Page 1)
    # ─────────────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("SUMMER TRAINING/INTERNSHIP REPORT\n(Term Aug-Dec 2026)\n\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    # Live Deployment Section
    p_live = doc.add_paragraph()
    p_live.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_live.paragraph_format.space_after = Pt(2)
    p_live.paragraph_format.left_indent = Inches(1.5)
    run_live = p_live.add_run("Live System Deployment Details:\n")
    run_live.font.name = 'Times New Roman'
    run_live.font.size = Pt(11.5)
    run_live.font.bold = True
    run_live.font.underline = True

    def add_cover_meta(label, url):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.5)
        p.paragraph_format.space_after = Pt(2)
        r_lbl = p.add_run(label)
        r_lbl.font.name = 'Times New Roman'
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(11)
        r_url = p.add_run(url)
        r_url.font.name = 'Times New Roman'
        r_url.font.size = Pt(11)
        
    add_cover_meta("GitHub Code Repository: ", "https://github.com/aryansingh1013/RepoVerseAi")
    add_cover_meta("Static Cloud Frontend: ", "https://aryansingh1013-repoverseai.static.hf.space")
    add_cover_meta("Docker API Web Service: ", "https://repoverseai.onrender.com (FastAPI backend)")

    doc.add_paragraph("\n\n")

    # Report Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(
        "A REPORT ON\n\n"
        "REPOVERSE AI: AN INTERACTIVE 3D GALAXY EXPLORER AND SEMANTIC RAG SEARCH PLATFORM FOR MODERN CODEBASES\n\n\n"
    )
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.font.bold = True

    # Submitted By
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(
        "SUBMITTED BY:\n\n"
        f"{name}\n"
        f"Registration Number: {reg_no}\n\n\n\n\n"
    )
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True

    # University Footer
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run(
        "School of Computer Science and Engineering\n"
        "Lovely Professional University, Phagwara, Punjab"
    )
    run_univ.font.name = 'Times New Roman'
    run_univ.font.size = Pt(13)
    run_univ.font.bold = True
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # DECLARATION (Page 2)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "DECLARATION", 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    decl_text = (
        f"I, {name}, hereby declare that the summer training/internship report entitled "
        "\"RepoVerse AI: An Interactive 3D Galaxy Explorer and Semantic RAG Search Platform for Modern Codebases\" "
        "submitted by me to the Department of Computer Science & Engineering, Lovely Professional University in "
        "partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in "
        "Computer Science and Engineering is an authentic record of my training carried out during the period "
        "from June 2026 to July 2026 under the supervision of , Assistant Professor.\n\n"
        "The matter embodied in this training report has not been submitted by me for the award of any "
        "other degree or diploma of this or any other University."
    )
    add_body_paragraph(doc, decl_text)
    
    doc.add_paragraph("\n\n\n\n")
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.paragraph_format.line_spacing = 1.15
    r_sig = p_sig.add_run(f"Date: 15 Jul '26\nPlace: Phagwara\n\n\n{name}\nRegistration Number: {reg_no}")
    r_sig.font.name = 'Times New Roman'
    r_sig.font.size = Pt(11.5)
    r_sig.font.bold = True
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # CERTIFICATE (Page 3)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "CERTIFICATE", 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    cert_text = (
        f"This is to certify that the Summer Training/Internship report entitled \"RepoVerse AI: An Interactive "
        "3D Galaxy Explorer and Semantic RAG Search Platform for Modern Codebases\" submitted by "
        f"{name} (Registration Number: {reg_no}) in partial fulfillment of the requirements for the award "
        "of the degree of Bachelor of Technology in Computer Science and Engineering, Lovely Professional "
        "University is a record of student's training work carried out under my supervision.\n\n"
        "To the best of my knowledge, the work presented here has reached the requisite standard for the "
        "submission of B.Tech summer training report."
    )
    add_body_paragraph(doc, cert_text)
    
    doc.add_paragraph("\n\n\n\n\n\n")
    p_cert_sig = doc.add_paragraph()
    p_cert_sig.paragraph_format.line_spacing = 1.25
    r_cert_sig = p_cert_sig.add_run(
        "___________________________\n"
        "[Rohit Bharti and Dipen Saini]\n"
        "Assistant Professor\n"
        "Department of Computer Science & Engineering\n\n\n\n\n\n"
        "___________________________\n"
        "Head of Department\n"
        "Department of Computer Science & Engineering"
    )
    r_cert_sig.font.name = 'Times New Roman'
    r_cert_sig.font.size = Pt(11.5)
    r_cert_sig.font.bold = True
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # ACKNOWLEDGEMENT (Page 4)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "ACKNOWLEDGEMENT", 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    ack_text = (
        "I would like to express my sincere gratitude and deep appreciation to my training supervisors, "
        "Rohit Bharti and Dipen Saini, Assistant Professor, Department of Computer Science & Engineering, for providing "
        "support, constructive advice, and guidance throughout my summer training curriculum and project execution. "
        "Their technical expertise and encouragement have been instrumental in the completion of this project.\n\n"
        "I also extend my heartfelt thanks to the Head of the Department and the administration of Lovely Professional "
        "University for providing the academic infrastructure and laboratories that facilitated this work. I am grateful "
        "to the open-source community behind FastAPI, LangChain, LangGraph, and React, whose libraries made this project "
        "modular, performant, and scalable.\n\n"
        "Finally, I would like to thank my parents, family, and peers for their continuous support, patience, and "
        "encouragement throughout this academic term."
    )
    add_body_paragraph(doc, ack_text)
    
    doc.add_paragraph("\n\n\n")
    p_ack_sub = doc.add_paragraph()
    p_ack_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ack_sub = p_ack_sub.add_run(f"{name}\nRegistration Number: {reg_no}")
    r_ack_sub.font.name = 'Times New Roman'
    r_ack_sub.font.size = Pt(11.5)
    r_ack_sub.font.bold = True
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS (Pages 5-6)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "TABLE OF CONTENTS", 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.35
    r_toc = toc_p.add_run(
        "Declaration\t\t\ti\n"
        "Certificate\t\t\tii\n"
        "Acknowledgement\t\t\tiii\n"
        "Table of Contents\t\t\tiv\n\n"
        "1. INTRODUCTION OF ORGANIZATION & ECOSYSTEM\t\t\t1\n"
        "   1.1 Overview of 3D Galaxy Visualization & Software Mapping\t\t\t1\n"
        "   1.2 Open-Source Ecosystem and Collaborative Development Framework\t\t\t2\n"
        "   1.3 Project Motivation: Solving Code Comprehension Bottlenecks\t\t\t3\n"
        "   1.4 Structural Goals and Scope of RepoVerse AI\t\t\t5\n\n"
        "2. SUMMER TRAINING COURSE & INTERNSHIP DETAILS\t\t\t6\n"
        "   2.1 Training Curriculum and Core Concepts Overview\t\t\t6\n"
        "   2.2 Week-by-Week Learning and Technical Milestones\t\t\t7\n"
        "       2.2.1 Week 1: Foundations of Asynchronous Scrapers and Walkers\t\t\t7\n"
        "       2.2.2 Week 2: AST Parsers, Code Slicing and Token Chunking\t\t\t9\n"
        "       2.2.3 Week 3: Vector Databases, Semantic Indexing and Search\t\t\t11\n"
        "       2.2.4 Week 4: FastAPI endpoints, WebSockets and Async Queue Management\t\t\t13\n"
        "       2.2.5 Week 5: LangGraph Agents, RAG pipelines and AI Skills Engine\t\t\t15\n"
        "       2.2.6 Week 6: React Three Fiber, 3D Graph layout and Docker deployments\t\t\t17\n\n"
        "3. REPOVERSE AI PROJECT DETAIL\t\t\t19\n"
        "   3.1 Problem Statement and Context\t\t\t19\n"
        "   3.2 Project Outcomes and Engineering Successes\t\t\t20\n"
        "   3.3 System Architecture and Processing Pipeline\t\t\t22\n"
        "   3.4 Detailed Technical Stack Mapping\t\t\t24\n\n"
        "4. SOURCE CODE AND SYSTEM SNAPSHOTS\t\t\t26\n"
        "   4.1 Key Implementation Snippets\t\t\t26\n"
        "       4.1.1 Subprocess Connection Managers\t\t\t26\n"
        "       4.1.2 Dynamic Skills Caching Layer\t\t\t28\n"
        "       4.1.3 Lazy Load Vector Embeddings\t\t\t30\n"
        "       4.1.4 Frontend Galaxy Interactions\t\t\t32\n"
        "   4.2 System Interface and Visual Snapshots\t\t\t34\n\n"
        "12. BIBLIOGRAPHY\t\t\t36\n"
    )
    r_toc.font.name = 'Times New Roman'
    r_toc.font.size = Pt(11.5)
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # CHAPTER 1 (Pages 7-10)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "1. INTRODUCTION OF ORGANIZATION & ECOSYSTEM", 1)
    
    add_heading_styled(doc, "1.1 OVERVIEW OF 3D GALAXY VISUALIZATION & SOFTWARE MAPPING", 2)
    p_1_1_a = (
        "Modern enterprise software codebases regularly comprise hundreds of directories and thousands of source code files, "
        "making comprehension, visualization, and onboarding extremely resource-intensive. Traditional text-based tools "
        "and flat IDE directories provide structural paths, but they fail to represent the logical density, topological coupling, "
        "and architectural hierarchies that exist between software modules. As a result, new developers face significant "
        "cognitive load when attempting to build a mental map of how files interact."
    )
    add_body_paragraph(doc, p_1_1_a)
    
    p_1_1_b = (
        "RepoVerse AI resolves this bottleneck by implementing a spatial mapping system that transforms abstract source code "
        "topologies into a highly engaging, interactive 3D space-themed galaxy. In this representation, the entire project "
        "corresponds to a galaxy, high-level directories correspond to stellar constellations, internal folders correspond to "
        "star systems, and individual source code files translate into planets. This visual paradigm allows developers to "
        "grasp code coupling and scale constraints at a single glance, turning directory trees into observable structures."
    )
    add_body_paragraph(doc, p_1_1_b)
    
    add_heading_styled(doc, "1.2 OPEN-SOURCE ECOSYSTEM AND COLLABORATIVE DEVELOPMENT FRAMEWORK", 2)
    p_1_2_a = (
        "This internship and research project was executed within the framework of an Open-Source Software Development Lab, "
        "simulating standard industrial Agile coding workflows. Collaborative software development is not merely about writing "
        "standalone functions; it requires establishing rigorous version control (using Git), code review policies, static "
        "analysis formatting check tools, and containerized build pipelines."
    )
    add_body_paragraph(doc, p_1_2_a)

    p_1_2_b = (
        "In this collaborative framework, we emphasized incremental integration. Each feature (like the dynamic skills engine "
        "and the lazy loading memory manager) was developed, locally debugged, and then pushed to a central repository "
        "hosted on GitHub. To support continuous integration, we configured build environments that test code compilation "
        "and package installations automatically. Adhering to these pipelines enforces standard python type checking, PEP-8 "
        "conformance, and modular class definitions."
    )
    add_body_paragraph(doc, p_1_2_b)

    add_heading_styled(doc, "1.3 PROJECT MOTIVATION: SOLVING CODE COMPREHENSION BOTTLENECKS", 2)
    p_1_3_a = (
        "The core motivation behind RepoVerse AI is the cognitive limitation developers face when onboarding onto "
        "complex codebases. Flat directory trees do not communicate hierarchy or complexity effectively. A developer "
        "cannot determine which modules are heavily imported (high incoming coupling) or which files contain thousands "
        "of lines of code, except by manual inspection. Visualizing these metrics as planetary orbits (where orbital speed, "
        "color, and size reflect complexity and incoming dependencies) bridges this gap."
    )
    add_body_paragraph(doc, p_1_3_a)

    p_1_3_b = (
        "Furthermore, modern AI chatbots often fall short because they lack grounding or exhaust API rate limits by reading "
        "raw files repeatedly. RepoVerse AI addresses this by providing a unified RAG (Retrieval-Augmented Generation) "
        "agent that is grounded on the vector index of the codebase. The motivation is to construct a self-contained, "
        "private, low-overhead RAG application that can run on consumer-grade hardware and free cloud services. By executing "
        "embeddings locally and calling Groq API completions with small, token-optimized context frames, we ensure "
        "maximum efficiency, high-speed execution, and strict data privacy."
    )
    add_body_paragraph(doc, p_1_3_b)

    add_heading_styled(doc, "1.4 STRUCTURAL GOALS AND SCOPE OF REPOVERSE AI", 2)
    p_1_4 = (
        "The scope of this project spans several engineering layers:\n"
        "1. Dynamic Codebase Parsing: The system must traverse any target folder recursively, clean node modules, and extract structural file-relationship hierarchies in milliseconds.\n"
        "2. Semantic Vector Indexing: Clean code chunks must be converted into high-dimensional vector embeddings and cached in a local vector database.\n"
        "3. 3D Web Graphics Interface: Render a responsive, glassmorphic UI utilizing Three.js and React Three Fiber to display codebase stars and orbiting planetary files.\n"
        "4. Token-Protected AI Agents: Create an AI Agent that answers technical queries about the codebase while citing source files, protected by disk-based caching to limit API token consumption."
    )
    add_body_paragraph(doc, p_1_4)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # CHAPTER 2 (Pages 11-18)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "2. SUMMER TRAINING COURSE & INTERNSHIP DETAILS", 1)
    
    add_heading_styled(doc, "2.1 TRAINING CURRICULUM AND CORE CONCEPTS OVERVIEW", 2)
    p_2_1 = (
        "The six-week training curriculum was meticulously designed to bridge the gap between academic python programming "
        "and professional full-stack AI engineering. The learning trajectory guided us through the entire lifecycle "
        "of modern AI applications: starting with Python's asynchronous libraries, passing through NLP parsing algorithms, "
        "vector mathematics, persistent databases, and moving to interactive 3D web rendering, reactive states, "
        "and multi-platform cloud deployments."
    )
    add_body_paragraph(doc, p_2_1)

    add_heading_styled(doc, "2.2 WEEK-BY-WEEK LEARNING AND TECHNICAL MILESTONES", 2)
    
    add_heading_styled(doc, "2.2.1 Week 1: Foundations of Asynchronous Scrapers and Walkers", 3)
    p_w1_a = (
        "During the first week, we focused on understanding Python's concurrency model. Traditional synchronous loops "
        "block the main thread while performing I/O operations (like reading files, waiting for subprocesses, or fetching API endpoints). "
        "We learned to implement Python's `asyncio` library, using asynchronous coroutines and tasks to handle hundreds "
        "of tasks concurrently."
    )
    add_body_paragraph(doc, p_w1_a)
    p_w1_b = (
        "We designed a directory walking utility that parses local folders. Unlike native `os.walk`, this async walker "
        "inspects folder depths and filters out blacklisted dependencies (like `node_modules` and `__pycache__`) in-place, "
        "drastically reducing file-system lookup latency. We also built an async Git crawler that searches for `.git` logs "
        "and reads commit lines asynchronously to construct history timelines."
    )
    add_body_paragraph(doc, p_w1_b)

    add_heading_styled(doc, "2.2.2 Week 2: AST Parsers, Code Slicing and Token Chunking", 3)
    p_w2_a = (
        "The second week centered on code analysis using Abstract Syntax Trees (ASTs). Rather than treating code as raw strings, "
        "ASTs represent code as hierarchical node trees. Python's built-in `ast` module allowed us to write code analyzers "
        "that parse class names, method signatures, argument types, global variables, and comments from python modules."
    )
    add_body_paragraph(doc, p_w2_a)
    p_w2_b = (
        "We also tackled the problem of code chunking. When indexing code for an LLM, we cannot feed it entire files, "
        "as it exceeds context sizes. We designed a semantic chunker that splits code files at functional boundaries (like "
        "class or def blocks) with 10% overlapping tokens to maintain contiguous context. For JS/TS files, we wrote regex-based "
        "scanners to extract module imports and define logical code blocks."
    )
    add_body_paragraph(doc, p_w2_b)

    add_heading_styled(doc, "2.2.3 Week 3: Vector Databases, Semantic Indexing and Search", 3)
    p_w3_a = (
        "Week three advanced our knowledge of vector mathematics and database indexing. We learned how sentence transformers "
        "convert text chunks into 384-dimensional or 768-dimensional float arrays (embeddings). In this vector space, semantically "
        "similar sentences are close to each other, evaluated using cosine similarity or Euclidean L2 distance."
    )
    add_body_paragraph(doc, p_w3_a)
    p_w3_b = (
        "We configured ChromaDB, an open-source vector database. We wrote collection wrappers to handle database persistence, "
        "metadata filtering, and vector additions. We also designed a Hybrid Retriever that combines Vector Search (for conceptual "
        "similarity) and BM25 lexical search (for exact keyword/variable name match) using Reciprocal Rank Fusion (RRF) algorithms."
    )
    add_body_paragraph(doc, p_w3_b)

    add_heading_styled(doc, "2.2.4 Week 4: FastAPI endpoints, WebSockets and Async Queue Management", 3)
    p_w4_a = (
        "In week four, we built the REST and WebSocket API infrastructure using FastAPI. FastAPI uses Python's ASGI specification "
        "to handle high-throughput request channels. We created routes to fetch project schemas, trigger folder indexing, "
        "and toggle visual themes."
    )
    add_body_paragraph(doc, p_w4_a)
    p_w4_b = (
        "We also created a WebSocket endpoint to stream AI chat completions. Standard HTTP requests wait for the entire "
        "LLM response before returning, which takes seconds. WebSockets allow the backend to stream tokens character-by-character "
        "in real-time. We implemented a WebSocket Connection Manager that coordinates active chat sessions, handles "
        "abrupt disconnects, and manages message histories."
    )
    add_body_paragraph(doc, p_w4_b)

    add_heading_styled(doc, "2.2.5 Week 5: LangGraph Agents, RAG pipelines and AI Skills Engine", 3)
    p_w5_a = (
        "Week five focused on Agentic workflows. We utilized LangGraph to define a stateful agent graph. A LangGraph agent "
        "can route its execution paths dynamically based on user prompts. For example, if a user asks a coding question, "
        "the agent routes to a RAG retriever; if the user asks a general question, it routes to a direct completion LLM."
    )
    add_body_paragraph(doc, p_w5_a)
    p_w5_b = (
        "We also built the AI Skills Engine. Instead of loading the entire codebase into the LLM context, we developed "
        "the hybrid dynamic analyzer: we use Python to perform static scans (extract dependencies, count file lines, search secrets) "
        "costing 0 tokens, and pass a highly compressed folder skeleton to the LLM (costing < 1100 tokens) to generate reports. "
        "These reports are saved in a local JSON cache to prevent duplicate LLM calls."
    )
    add_body_paragraph(doc, p_w5_b)

    add_heading_styled(doc, "2.2.6 Week 6: React Three Fiber, 3D Graph layout and Docker deployments", 3)
    p_w6_a = (
        "In the final week, we developed the frontend interface and handled DevOps. We used React Three Fiber to write "
        "3D components. We mapped the folder hierarchy into a 3D gravity-based layout: stars at the center, planets orbiting "
        "stars, and moons orbiting planets, rendering them with customizable shader materials and lighting."
    )
    add_body_paragraph(doc, p_w6_a)
    p_w6_b = (
        "For deployment, we wrote a multi-stage Dockerfile and deployed the backend to Render (optimizing the code to "
        "lazy-load heavy PyTorch embeddings so that it stays under 512MB RAM). We also compiled the React frontend with "
        "production URLs and hosted it statically on Hugging Face Spaces using HashRouter to prevent iframe sandboxing errors."
    )
    add_body_paragraph(doc, p_w6_b)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # CHAPTER 3 (Pages 19-25)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "3. REPOVERSE AI PROJECT DETAIL", 1)
    
    add_heading_styled(doc, "3.1 PROBLEM STATEMENT AND CONTEXT", 2)
    p_3_1_a = (
        "Onboarding a software developer onto a large legacy repository is a slow, error-prone, and expensive process. "
        "Flat file structures do not show how components interact, which directories are heavily modified, or what APIs "
        "serve as system inputs. Standard IDE search is limited to exact keyword lookups and doesn't support conceptual queries."
    )
    add_body_paragraph(doc, p_3_1_a)

    p_3_1_b = (
        "Furthermore, using plain LLM agents directly on codebases introduces engineering challenges: they quickly hit token limit "
        "thresholds, hallucinate function signatures, and leak proprietary code to external APIs. There is a critical need "
        "for a visual code explorer that combines structural representation (how files relate) and semantic intelligence (AI chat) "
        "in a performant, cost-controlled, and private deployment."
    )
    add_body_paragraph(doc, p_3_1_b)

    add_heading_styled(doc, "3.2 PROJECT OUTCOMES AND ENGINEERING SUCCESSES", 2)
    p_3_2 = (
        "RepoVerse AI successfully met these requirements. The major project outcomes include:\n"
        "1. Interactive 3D Galaxy Interface: Folders render as stellar constellations and files render as orbiting planets. Hovering displays file metrics (line counts, functions), and clicking opens the code viewer.\n"
        "2. Zero-Token Static Profilers: Built AST-based dependency explorers and git timeline crawlers that analyze the local repository statically, utilizing 0 LLM API credits.\n"
        "3. Disk-based JSON Cache: Cached AI skill analysis results in `.repoverse/skills_cache.json` so repeat reports load instantly and cost 0 tokens.\n"
        "4. Clean Memory Footprint: Modified imports and embedding managers to lazy-load PyTorch models. This keeps startup RAM under 80MB, allowing free tier cloud hosting on Render without Out-Of-Memory crashes.\n"
        "5. Iframe Sandbox Compatibility: Implemented HashRouter in the React app, allowing seamless 3D rendering inside Hugging Face sandboxed iframes without security blocks."
    )
    add_body_paragraph(doc, p_3_2)

    add_heading_styled(doc, "3.3 SYSTEM ARCHITECTURE AND PROCESSING PIPELINE", 2)
    p_3_3 = (
        "The system architecture follows a decoupled model. The React client handles 3D rendering and state controls, while "
        "the FastAPI backend handles directory parsing, vector database storage, and agent execution. When a user indexes a repo:\n"
        "1. The AST Parser scans the directories recursively.\n"
        "2. The file nodes are built into a hierarchical JSON schema.\n"
        "3. The Embeddings Manager converts semantic code chunks into float vectors.\n"
        "4. ChromaDB indexes these vectors for hybrid retrieval.\n"
        "5. User queries are routed through a LangGraph agent that checks the vector database, formats contextual prompts, and streams responses."
    )
    add_body_paragraph(doc, p_3_3)

    add_heading_styled(doc, "3.4 DETAILED TECHNICAL STACK MAPPING", 2)
    
    # Custom styled table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology Layer'
    hdr_cells[1].text = 'Framework/Library'
    hdr_cells[2].text = 'Role in System Architecture'
    
    for cell in hdr_cells:
        set_cell_background(cell, "4B5563") # Dark gray header
        set_cell_margins(cell, top=120, bottom=120)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10.5)

    tech_stack = [
        ("Web Framework", "FastAPI / Uvicorn", "Serves REST endpoints, runs WebSocket pipelines, and routes static files."),
        ("Vector DB", "ChromaDB", "Persists local vectors, tracks document IDs, and runs semantic search queries."),
        ("Orchestrator", "LangGraph / LangChain", "Orchestrates stateful agent paths, structures prompt templates, and runs tools."),
        ("3D Graphics", "React Three Fiber / Three.js", "Compiles webGL shaders, renders planets, orbit lines, and lighting."),
        ("Embeddings", "BAAI/bge-small-en-v1.5", "Embeds text chunks locally using SentenceTransformer models."),
        ("LLM API Provider", "Groq Cloud (Llama 3.3)", "Executes fast, grounded completions based on code context."),
        ("UI Components", "React (v18) / TailwindCSS", "Renders lists, charts, and skills layouts.")
    ]

    for layer, lib, role in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = lib
        row_cells[2].text = role
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80)
            set_cell_background(cell, "F9FAFB")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # CHAPTER 4 (Pages 26-35)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "4. SOURCE CODE AND SYSTEM SNAPSHOTS", 1)
    
    add_heading_styled(doc, "4.1 KEY IMPLEMENTATION SNIPPETS", 2)
    p_snip = (
        "This chapter lists the core Python and React files implemented to support dynamic skills caching, "
        "memory optimization, and sandboxed rendering."
    )
    add_body_paragraph(doc, p_snip)

    add_heading_styled(doc, "4.1.1 Subprocess Connection Managers (connection_manager.py)", 3)
    p_code1_desc = (
        "The following snippet handles the background loading of Model Context Protocol (MCP) servers. "
        "It manages subprocesses asynchronously to prevent connection blocking on startup:"
    )
    add_body_paragraph(doc, p_code1_desc)

    code1 = (
        "class MCPConnectionManager:\n"
        "    def __init__(self, registry: MCPServerRegistry):\n"
        "        self.registry = registry\n"
        "        self.active_servers: Dict[str, MCPSubprocessServer] = {}\n"
        "        self.execution_counts: Dict[str, int] = {}\n"
        "        self.error_logs: Dict[str, List[str]] = {}\n\n"
        "    def connect_server(self, name: str, command: List[str], env: Optional[Dict[str, str]] = None) -> bool:\n"
        "        self.disconnect_server(name)\n"
        "        timeout = mcp_settings.get(\"connection_timeout\", 30)\n"
        "        server = MCPSubprocessServer(name, command, env=env, timeout=timeout)\n"
        "        start_time = time.perf_counter()\n"
        "        success = server.initialize()\n"
        "        latency = (time.perf_counter() - start_time) * 1000.0\n"
        "        if success:\n"
        "            self.active_servers[name] = server\n"
        "            self.registry.register(name, server)\n"
        "            return True\n"
        "        return False"
    )
    add_code_block(doc, code1)

    add_heading_styled(doc, "4.1.2 Dynamic Skills Caching Layer (cache.py)", 3)
    p_code2_desc = (
        "To avoid repeated LLM calls during exploration, report outputs are stored under the workspace cache. "
        "This completely eliminates duplicate token consumption:"
    )
    add_body_paragraph(doc, p_code2_desc)

    code2 = (
        "def get_cached(slug: str, workspace_dir: str) -> Optional[Dict[str, Any]]:\n"
        "    cache = _load_cache(workspace_dir)\n"
        "    key = f\"{slug}:{_workspace_hash(workspace_dir)}\"\n"
        "    result = cache.get(key)\n"
        "    if result is not None:\n"
        "        print(f\"SkillCache: Cache HIT for '{slug}' — 0 tokens used.\")\n"
        "    return result\n\n"
        "def set_cached(slug: str, workspace_dir: str, result: Dict[str, Any]) -> None:\n"
        "    cache = _load_cache(workspace_dir)\n"
        "    key = f\"{slug}:{_workspace_hash(workspace_dir)}\"\n"
        "    cache[key] = result\n"
        "    _save_cache(workspace_dir, cache)"
    )
    add_code_block(doc, code2)

    add_heading_styled(doc, "4.1.3 Lazy Load Vector Embeddings (vector_store.py)", 3)
    p_code3_desc = (
        "To stay within Render's free 512MB RAM limit, we refactored the embeddings class. The heavy "
        "SentenceTransformer import statement is moved inside the property, meaning it is only loaded on use:"
    )
    add_body_paragraph(doc, p_code3_desc)

    code3 = (
        "class EmbeddingsManager:\n"
        "    def __init__(self):\n"
        "        self.model_name = settings.EMBEDDING_MODEL\n"
        "        self._model = None\n\n"
        "    @property\n"
        "    def model(self):\n"
        "        if self._model is None:\n"
        "            try:\n"
        "                print(f\"Lazy loading embedding model {self.model_name}...\")\n"
        "                from sentence_transformers import SentenceTransformer\n"
        "                self._model = SentenceTransformer(self.model_name)\n"
        "            except Exception as e:\n"
        "                from sentence_transformers import SentenceTransformer\n"
        "                self._model = SentenceTransformer(settings.EMBEDDING_MODEL_FALLBACK)\n"
        "        return self._model"
    )
    add_code_block(doc, code3)

    add_heading_styled(doc, "4.1.4 Frontend Galaxy Routing (App.tsx)", 3)
    p_code4_desc = (
        "To prevent sandbox iframe blocking on Hugging Face Spaces, the routing is set to HashRouter, "
        "bypassing HTML5 History API permissions:"
    )
    add_body_paragraph(doc, p_code4_desc)

    code4 = (
        "import { HashRouter, Routes, Route } from \"react-router-dom\";\n"
        "import { NavigationProvider } from \"@/hooks/useNavigation\";\n"
        "import { LandingPage } from \"@/pages/LandingPage\";\n"
        "import { IndexingScreen } from \"@/pages/IndexingScreen\";\n"
        "import { WorkspacePage } from \"@/pages/WorkspacePage\";\n\n"
        "export default function App() {\n"
        "  return (\n"
        "    <HashRouter>\n"
        "      <NavigationProvider>\n"
        "        <Routes>\n"
        "          <Route path=\"/\" element={<LandingPage />} />\n"
        "          <Route path=\"/indexing\" element={<IndexingScreen />} />\n"
        "          <Route path=\"/workspace\" element={<WorkspacePage />} />\n"
        "        </Routes>\n"
        "      </NavigationProvider>\n"
        "    </HashRouter>\n"
        "  );\n"
        "}"
    )
    add_code_block(doc, code4)

    add_heading_styled(doc, "4.2 SYSTEM INTERFACE AND VISUAL SNAPSHOTS", 2)
    p_snapshot_desc = (
        "Below are placeholders representing the core layout components. [Refer to the screenshots in the docs folder for final visuals]:\n\n"
        "• [Insert Figure 1: 3D Galaxy Explorer View] — Displays files as rotating planetary spheres orbiting constellation stars, with depth color mappings.\n\n"
        "• [Insert Figure 2: AI Mission Control Panel] — Side-pane chat window showing code questions, responses, and clickable path citations.\n\n"
        "• [Insert Figure 3: Skills Console overlay] — Displays custom analysis reports (Overview, Security audits, health scorecards) with the explicit Execute button."
    )
    add_body_paragraph(doc, p_snapshot_desc)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BIBLIOGRAPHY (Page 36)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading_styled(doc, "12. BIBLIOGRAPHY", 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    p_bib = doc.add_paragraph()
    p_bib.paragraph_format.line_spacing = 1.3
    r_bib = p_bib.add_run(
        "• FastAPI Framework Reference. https://fastapi.tiangolo.com/ - "
        "Consulted for ASGI standard specifications, request routing, and WebSockets logic.\n\n"
        "• LangGraph State Management Guide. https://langchain-ai.github.io/langgraph/ - "
        "Referenced for building routing nodes, conditional branches, and conversation histories.\n\n"
        "• ChromaDB Persistent Storage. https://www.trychroma.com/ - "
        "Consulted for collection parameters, sqlite indexing, and embedding queries.\n\n"
        "• React Three Fiber 3D Graphics. https://docs.pmnd.rs/react-three-fiber - "
        "Referenced for WebGL canvas setups, point controls, and planetary coordinate calculations.\n\n"
        "• Hugging Face Spaces Docker and Static Hosting. https://huggingface.co/docs/hub/spaces - "
        "Consulted for iframe sandbox definitions, folder structures, and relative assets configurations.\n\n"
        "• Groq LLM Speeds and speculative decoding metrics. https://console.groq.com/docs - "
        "Referenced for temperature tuning, token limitations, and API completion performance."
    )
    r_bib.font.name = 'Times New Roman'
    r_bib.font.size = Pt(11.5)

    doc.save(output_path)
    print(f"Massive report successfully saved to {output_path}!")

if __name__ == "__main__":
    out = os.path.abspath(os.path.join(os.path.dirname(__file__), "RepoVerse_AI_Summer_Training_Report.docx"))
    create_massive_report(out, "Aryan Singh", "12408478")
